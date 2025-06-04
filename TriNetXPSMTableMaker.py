# streamlit_app.py
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

st.title("TriNetX Table Formatter for Journal Submission")

uploaded_file = st.file_uploader("Upload your TriNetX CSV file", type="csv")

def extract_clean_table(df):
    for i, row in df.iterrows():
        if "Characteristic" in str(row[0]):
            data_start_index = i
            break
    df_clean = df.iloc[data_start_index:].reset_index(drop=True)
    df_clean.columns = df_clean.iloc[0]
    df_clean = df_clean[1:].reset_index(drop=True)
    return df_clean

def generate_word(df, font_size=10, align="left"):
    doc = Document()
    doc.add_heading('Formatted Table', level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.font.size = Pt(font_size)
        hdr_cells[i].paragraphs[0].alignment = {"left": 0, "center": 1, "right": 2}.get(align, 0)

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = str(cell)
            run = row_cells[i].paragraphs[0].runs[0]
            run.font.size = Pt(font_size)
            row_cells[i].paragraphs[0].alignment = {"left": 0, "center": 1, "right": 2}.get(align, 0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_pdf(df, font_size=8, align="L"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=font_size)
    col_width = pdf.w / (len(df.columns) + 1)
    row_height = font_size + 2

    for col in df.columns:
        pdf.cell(col_width, row_height, str(col), border=1, align=align)
    pdf.ln(row_height)

    for _, row in df.iterrows():
        for item in row:
            pdf.cell(col_width, row_height, str(item), border=1, align=align)
        pdf.ln(row_height)

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file)
    cleaned_df = extract_clean_table(df)

    st.sidebar.header("Display Options")
    show_gridlines = st.sidebar.checkbox("Show Gridlines", value=True)
    bold_headers = st.sidebar.checkbox("Bold Headers", value=True)
    decimals = st.sidebar.slider("Decimal Places", 0, 5, 2)

    export_font_size = st.sidebar.slider("Word Font Size", 6, 14, 10)
    export_align = st.sidebar.selectbox("Word Text Alignment", ["left", "center", "right"])

    pdf_font_size = st.sidebar.slider("PDF Font Size", 6, 14, 8)
    pdf_align = st.sidebar.selectbox("PDF Text Alignment", ["L", "C", "R"])

    df_display = cleaned_df.copy()
    for col in df_display.columns:
        try:
            df_display[col] = df_display[col].astype(float).round(decimals)
        except:
            pass

    st.markdown("### Preview of Formatted Table")
    if bold_headers:
        st.write(f"<style>thead th {{ font-weight: bold; }}</style>", unsafe_allow_html=True)

    if show_gridlines:
        styled = df_display.style.set_table_styles(
            [{'selector': 'td, th', 'props': [('border', '1px solid black')]}]
        ).set_properties(**{'border': '1px solid black'})
        st.write(styled)
    else:
        st.write(df_display)

    csv = df_display.to_csv(index=False).encode('utf-8')
    st.download_button("Download CSV", csv, "formatted_table.csv", "text/csv")

    word_file = generate_word(df_display, font_size=export_font_size, align=export_align)
    st.download_button("Download as Word", word_file, "formatted_table.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    pdf_file = generate_pdf(df_display, font_size=pdf_font_size, align=pdf_align)
    st.download_button("Download as PDF", pdf_file, "formatted_table.pdf", "application/pdf")
